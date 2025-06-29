from flask import Flask, request, render_template, session, redirect, url_for, jsonify, Response
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
from pymongo import MongoClient
from bson.objectid import ObjectId
from pymilvus import connections, Collection, FieldSchema, CollectionSchema, DataType, utility
from werkzeug.security import generate_password_hash, check_password_hash
import requests
import os
import uuid
import io
from dotenv import load_dotenv
from datetime import datetime
import torch
import hashlib
from bson.binary import Binary
import docx  # Added for DOCX support
import time
from functools import wraps
import json

# Check environment setup
def check_env_file():
    """Check if the .env file exists and print debug information."""
    env_file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    
    if os.path.exists(env_file_path):
        print(f"[INFO] Found .env file")
        try:
            with open(env_file_path, 'r') as f:
                lines = f.readlines()
            
            has_api_key = False
            for line in lines:
                if line.strip().startswith('OPENROUTER_API_KEY='):
                    has_api_key = True
                    break
                    
            if not has_api_key:
                print("[WARNING] OPENROUTER_API_KEY not found in .env file")
        except Exception as e:
            print(f"[ERROR] Error reading .env file: {str(e)}")
    else:
        print("[WARNING] No .env file found. Run setup_env.py to create one.")

# Load environment variables
load_dotenv()
env_api_key = os.getenv("OPENAI_API_KEY")
hardcoded_api_key = ""

# Milvus connection settings
MILVUS_HOST = os.getenv("MILVUS_HOST", "localhost")
MILVUS_PORT = os.getenv("MILVUS_PORT", "19530")

# Try to get API key from environment first, fall back to hardcoded if needed
if env_api_key:
    OPENAI_API_KEY = env_api_key
else:
    OPENAI_API_KEY = hardcoded_api_key

# Print API key status at startup
if OPENAI_API_KEY:
    masked_key = OPENAI_API_KEY[:4] + "..." + OPENAI_API_KEY[-4:] if len(OPENAI_API_KEY) > 8 else "***"
    print(f"[INFO] OpenAI API key loaded: {masked_key}")
else:
    print("[WARNING] No OpenAI API key found - API functionality will be limited")

# Disable fallback responses
USE_FALLBACK_RESPONSES = False

# Detect GPU
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# Load embedding model
embedder = SentenceTransformer("all-MiniLM-L6-v2", device=DEVICE)
EMBED_DIM = 384

# Flask app
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 1 * 1024 * 1024 * 1024  # 1 GB
app.secret_key = os.urandom(24)

# Connect to MongoDB
client = MongoClient("mongodb://3.110.121.81:27017/")
db = client['assistant_db']
users = db['users']
conversations = db['conversations']

# Connect to Milvus
COLLECTION_NAME = "document_chunks"

def connect_to_milvus():
    """Connect to Milvus with retry logic"""
    max_retries = 3
    retry_delay = 2  # seconds
    
    for attempt in range(max_retries):
        try:
            print(f"[DEBUG] Attempting to connect to Milvus (attempt {attempt + 1}/{max_retries})")
            connections.connect(
                alias="default",
                host=MILVUS_HOST,
                port=MILVUS_PORT
            )
            print("[INFO] Successfully connected to Milvus")
            return True
        except Exception as e:
            print(f"[ERROR] Failed to connect to Milvus (attempt {attempt + 1}/{max_retries}): {str(e)}")
            if attempt < max_retries - 1:
                print(f"[INFO] Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
            else:
                print("[ERROR] All connection attempts failed")
                return False
    
    return False

def init_milvus():
    """Initialize Milvus collection with proper schema"""
    try:
        # First ensure connection is established
        if not connect_to_milvus():
            raise Exception("Failed to connect to Milvus")
            
        if utility.has_collection(COLLECTION_NAME):
            print(f"[INFO] Collection {COLLECTION_NAME} already exists")
            return
            
        print(f"[INFO] Creating new collection {COLLECTION_NAME}")
        fields = [
            FieldSchema(name="id", dtype=DataType.VARCHAR, is_primary=True, auto_id=False, max_length=36),
            FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=2200),  # Increased from 2000 for safety
            FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=EMBED_DIM),
            FieldSchema(name="file_hash", dtype=DataType.VARCHAR, max_length=32),
            FieldSchema(name="file_name", dtype=DataType.VARCHAR, max_length=255),
            FieldSchema(name="username", dtype=DataType.VARCHAR, max_length=100)  # Add username field
        ]
        schema = CollectionSchema(fields, description="Document chunks with embeddings and user ownership")
        collection = Collection(name=COLLECTION_NAME, schema=schema)
        
        # Create index only for vector field
        collection.create_index(
            field_name="embedding",
            index_params={"metric_type": "L2", "index_type": "IVF_FLAT", "params": {"nlist": 128}}
        )
        collection.load()
        print(f"[INFO] Collection {COLLECTION_NAME} created and loaded successfully")
    except Exception as e:
        print(f"[ERROR] Critical error initializing Milvus: {str(e)}")
        print("[WARNING] Application will start but vector search functionality will be limited")
        raise

# Initialize Milvus with error handling
try:
    init_milvus()
except Exception as e:
    print(f"[ERROR] Critical error initializing Milvus: {str(e)}")
    print("[WARNING] Application will start but vector search functionality will be limited")

def with_milvus_connection(func):
    """Decorator to ensure Milvus connection before executing operations"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        if not connect_to_milvus():
            raise Exception("Failed to connect to Milvus")
        return func(*args, **kwargs)
    return wrapper

@with_milvus_connection
def get_milvus_collection():
    """Get Milvus collection with proper error handling"""
    try:
        collection = Collection(COLLECTION_NAME)
        collection.load()
        return collection
    except Exception as e:
        print(f"[ERROR] Failed to get collection: {str(e)}")
        raise

# Allowed file extensions
ALLOWED_FILE_EXTENSIONS = {'pdf', 'docx'}

# File utility functions
def is_allowed_file(filename):
    """Check if filename is an allowed file type"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_FILE_EXTENSIONS

def extract_pdf_text(file):
    """Extract text from PDF file"""
    print(f"[DEBUG] Starting PDF extraction for file: {file.filename}")
    reader = PdfReader(file)
    print(f"[DEBUG] PDF has {len(reader.pages)} pages")
    
    extracted_text = ""
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            extracted_text += page_text + "\n"
            print(f"[DEBUG] Extracted {len(page_text)} characters from page {i+1}")
        else:
            print(f"[DEBUG] No text extracted from page {i+1}")
    
    print(f"[DEBUG] Total extracted text: {len(extracted_text)} characters")
    return extracted_text

def extract_docx_text(file):
    """Extract text from DOCX file"""
    print(f"[DEBUG] Starting DOCX extraction for file: {file.filename}")
    doc = docx.Document(file)
    
    extracted_text = ""
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            extracted_text += paragraph.text + "\n"
            print(f"[DEBUG] Extracted paragraph {i+1}: {len(paragraph.text)} characters")
    
    # Also extract text from tables
    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                if cell.text:
                    extracted_text += cell.text + "\n"
                    print(f"[DEBUG] Extracted text from table {i+1}, row {j+1}, cell {k+1}")
    
    print(f"[DEBUG] Total extracted text: {len(extracted_text)} characters")
    return extracted_text

def extract_text_from_file(file):
    """Extract text from either PDF or DOCX file based on extension"""
    filename = file.filename.lower()
    if filename.endswith('.pdf'):
        return extract_pdf_text(file)
    elif filename.endswith('.docx'):
        return extract_docx_text(file)
    return ""

def chunk_text(text, chunk_size=1700, overlap=200):
    """Split text into chunks with overlap"""
    print(f"[DEBUG] Starting text chunking - total text length: {len(text)} chars")
    print(f"[DEBUG] Using chunk_size={chunk_size}, overlap={overlap}")
    
    chunks, start = [], 0
    chunk_count = 0
    
    while start < len(text):
        # Extract chunk but ensure we don't exceed the safe limit
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]
        
        # Double-check length and truncate if needed
        if len(chunk) > 1900:  # Keep a safety buffer below the 2000 limit
            print(f"[WARNING] Chunk exceeded safety limit: {len(chunk)} chars. Truncating.")
            chunk = chunk[:1900]
        
        chunks.append(chunk)
        chunk_count += 1
        print(f"[DEBUG] Created chunk #{chunk_count}: start={start}, length={len(chunk)}")
        
        # Move start position for next chunk, accounting for overlap
        start += chunk_size - overlap
        
    print(f"[DEBUG] Chunking complete. Created {len(chunks)} chunks")
    
    # Print sample of first and last chunk for verification
    if chunks:
        print(f"[DEBUG] First chunk sample: {chunks[0][:100]}...")
        print(f"[DEBUG] Last chunk sample: {chunks[-1][:100]}...")
    
    return chunks

# Milvus storage
def batch_insert(data, batch_size=1000):
    """
    Insert data into Milvus in batches
    
    Args:
        data: List of dictionaries containing the data to insert
        batch_size: Number of items to insert in each batch
    """
    print(f"[DEBUG] Batch insert starting with {len(data)} items, batch_size={batch_size}")
    collection = Collection(COLLECTION_NAME)
    
    # Validate that all entries have required fields
    for item in data:
        if "username" not in item:
            raise ValueError("All entries must have a username field")
    
    for i in range(0, len(data), batch_size):
        batch = data[i:i + batch_size]
        print(f"[DEBUG] Processing batch #{i//batch_size + 1}: items {i} to {min(i+batch_size, len(data))}")
        
        try:
            print(f"[DEBUG] Preparing batch data arrays")
            ids = [item["id"] for item in batch]
            texts = [item["text"] for item in batch]
            embeddings = [item["embedding"] for item in batch]
            file_hashes = [item["file_hash"] for item in batch]
            file_names = [item["file_name"] for item in batch]
            usernames = [item["username"] for item in batch]  # Extract usernames
            
            print(f"[DEBUG] Inserting batch into Milvus collection: {COLLECTION_NAME}")
            collection.insert([ids, texts, embeddings, file_hashes, file_names, usernames])  # Include usernames in insert
            print(f"[DEBUG] Batch insertion successful")
            
            print(f"[DEBUG] Flushing collection")
            collection.flush()
            print(f"[DEBUG] Loading collection")
            collection.load()
            print(f"[DEBUG] Batch #{i//batch_size + 1} complete")
        except Exception as e:
            error_str = str(e)
            print(f"[ERROR] Exception during batch insert: {error_str}")
            
            # Parse the error message to identify problematic rows
            if "length of varchar field text exceeds max length" in error_str:
                try:
                    # Extract the row number from the error message
                    import re
                    row_match = re.search(r"row number: (\d+)", error_str)
                    if row_match:
                        row_number = int(row_match.group(1))
                        # Convert to zero-based index in our batch
                        row_index = row_number - 1
                        if 0 <= row_index < len(batch):
                            print(f"[INFO] Problem identified in row {row_number} (index {row_index})")
                            print(f"[INFO] Text length: {len(texts[row_index])}")
                            
                            # Fix just the problematic text by making it significantly shorter
                            texts[row_index] = texts[row_index][:1800]
                            print(f"[INFO] Truncated text to length: {len(texts[row_index])}")
                            
                            # Try insertion again with the fixed text
                            print(f"[INFO] Retrying insertion with truncated text")
                            collection.insert([ids, texts, embeddings, file_hashes, file_names, usernames])  # Include usernames in retry
                            print(f"[INFO] Retry successful")
                            
                            # Proceed with flush and load
                            collection.flush()
                            collection.load()
                            print(f"[DEBUG] Batch #{i//batch_size + 1} complete after retry")
                            continue
                    
                    # If we couldn't identify the specific row, do a more aggressive truncation on all rows
                    print("[WARNING] Could not identify specific problematic row. Truncating all texts to 1800 chars.")
                    texts = [t[:1800] for t in texts]
                    collection.insert([ids, texts, embeddings, file_hashes, file_names, usernames])  # Include usernames in retry
                    print(f"[INFO] Retry with all texts truncated successful")
                    
                    # Proceed with flush and load
                    collection.flush()
                    collection.load()
                    print(f"[DEBUG] Batch #{i//batch_size + 1} complete after retry")
                    
                except Exception as retry_error:
                    print(f"[ERROR] Retry also failed: {str(retry_error)}")
                    # Try one last time with extremely conservative truncation
                    try:
                        texts = [t[:1500] for t in texts]
                        print("[INFO] Last attempt with all texts truncated to 1500 chars")
                        collection.insert([ids, texts, embeddings, file_hashes, file_names, usernames])  # Include usernames in final retry
                        print(f"[INFO] Final retry successful")
                        collection.flush()
                    except Exception as final_error:
                        print(f"[ERROR] Final retry also failed: {str(final_error)}")
                        raise
            else:
                # Not a text length issue, re-raise
                raise

def get_user_documents(username):
    """Get all documents for a specific user"""
    if not username:
        raise ValueError("Username is required")
        
    collection = Collection(COLLECTION_NAME)
    collection.load()
    
    try:
        # Query only documents belonging to this user
        results = collection.query(
            expr=f'username == "{username}"',
            output_fields=["id", "text", "file_name", "file_hash"]
        )
        return results
    except Exception as e:
        print(f"[ERROR] Failed to get user documents: {str(e)}")
        return []

def delete_user_documents(username):
    """Delete all documents for a specific user"""
    if not username:
        raise ValueError("Username is required")
        
    collection = Collection(COLLECTION_NAME)
    collection.load()
    
    try:
        # Delete only documents belonging to this user
        collection.delete(f'username == "{username}"')
        print(f"[INFO] Deleted all documents for user: {username}")
    except Exception as e:
        print(f"[ERROR] Failed to delete user documents: {str(e)}")
        raise

def verify_document_ownership(document_id, username):
    """Verify that a document belongs to a specific user"""
    if not username or not document_id:
        return False
        
    collection = Collection(COLLECTION_NAME)
    collection.load()
    
    try:
        # Query the document and check ownership
        results = collection.query(
            expr=f'id == "{document_id}" and username == "{username}"',
            output_fields=["id"]
        )
        return len(results) > 0
    except Exception as e:
        print(f"[ERROR] Failed to verify document ownership: {str(e)}")
        return False

def calculate_file_hash(file):
    """Calculate hash of file content"""
    file.seek(0)
    return hashlib.md5(file.read()).hexdigest()

def store_chunks_in_milvus(chunks, username):
    """
    Store document chunks in Milvus with user ownership
    
    Args:
        chunks: List of document chunks to store
        username: Username of the user who owns these chunks
    """
    print(f"[DEBUG] Beginning to encode {len(chunks)} chunks with embedder")
    print(f"[DEBUG] Embedder model: {embedder}")
    print(f"[DEBUG] Using device: {DEVICE}")
    print(f"[DEBUG] Storing chunks for user: {username}")
    
    try:
        # Calculate total chunks and initialize progress
        total_chunks = len(chunks)
        processed_chunks = 0
        failed_chunks = 0
        
        # Process chunks in batches
        batch_size = 32
        data = []
        
        for i in range(0, total_chunks, batch_size):
            batch = chunks[i:i + batch_size]
            
            try:
                # Encode the batch
                vectors = embedder.encode([chunk["text"] for chunk in batch], convert_to_numpy=True, show_progress_bar=False)
                
                for j, chunk in enumerate(batch):
                    # Truncate text to safe limit here, before adding to data
                    text = chunk["text"]
                    if len(text) > 1900:  # Preemptively truncate
                        print(f"[WARNING] Preemptively truncating text from {len(text)} to 1900 chars")
                        text = text[:1900]
                        
                    data.append({
                        "id": chunk["id"],
                        "text": text,
                        "embedding": vectors[j].tolist(),
                        "file_hash": chunk["file_hash"],
                        "file_name": chunk["file_name"],
                        "username": username  # Add username to each chunk
                    })
                
                processed_chunks += len(batch)
                progress = (processed_chunks / total_chunks) * 100
                print(f"[PROGRESS] {progress:.1f}% - Processed {processed_chunks}/{total_chunks} chunks")
                
                # Insert batch into Milvus
                if data:
                    try:
                        batch_insert(data)
                        data = []  # Clear data after successful insert
                    except Exception as e:
                        # If insertion fails, log it but continue with next batch
                        failed_chunks += len(batch)
                        print(f"[ERROR] Failed to insert batch: {str(e)}")
                        # We'll yield an error message later but continue processing
                        data = []  # Clear data to proceed with next batch
                
                # Send progress update to frontend
                yield f"data: {progress:.1f}\n\n"
                
            except Exception as batch_error:
                # Handle errors in a batch but continue with others
                failed_chunks += len(batch)
                print(f"[ERROR] Error processing batch: {str(batch_error)}")
                # We'll yield an error message later but continue processing
        
        print(f"[DEBUG] Milvus insertion complete")
        
        # Report final status
        if failed_chunks > 0:
            failure_percent = (failed_chunks / total_chunks) * 100
            message = f"Processing complete with errors. {failed_chunks} chunks ({failure_percent:.1f}%) failed to process."
            print(f"[WARNING] {message}")
            yield f"data: {100.0}\n\n"
            yield f"data: warning: {message}\n\n"
        else:
            yield f"data: 100.0\n\n"
        
    except Exception as e:
        print(f"[ERROR] Exception in store_chunks_in_milvus: {str(e)}")
        yield f"data: error: {str(e)}\n\n"
        # Don't re-raise - allow partial progress to be reported to the user

# Keyword-based search in Milvus
def keyword_search(query, top_k=20, username=None):
    """
    Search for chunks containing keywords from the query
    
    Args:
        query: The search query
        top_k: Number of results to return
        username: Optional username to filter results by
    """
    print(f"[DEBUG] Starting keyword search for query: '{query}'")
    print(f"[DEBUG] Requested top_k={top_k}, username={username}")
    
    try:
        # Process query to extract keywords
        # Remove common words and punctuation
        import re
        from collections import Counter
        
        # Clean query and extract keywords
        cleaned_query = re.sub(r'[^\w\s]', '', query.lower())
        words = cleaned_query.split()
        
        # Filter out common words (simple approach)
        common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'with', 'is', 'are', 'was', 'were'}
        keywords = [word for word in words if word not in common_words and len(word) > 2]
        
        if not keywords:
            print(f"[DEBUG] No meaningful keywords extracted from query")
            return []
            
        print(f"[DEBUG] Extracted keywords: {keywords}")
        
        # Connect to Milvus
        collection = Collection(COLLECTION_NAME)
        collection.load()
        
        # Create keyword search expression
        # For each keyword, search for chunks containing that keyword
        keyword_conditions = []
        for keyword in keywords:
            # Use case-insensitive search with LIKE operator
            keyword_conditions.append(f"text like '%{keyword}%'")
        
        # Combine conditions with OR (any keyword match)
        expr = " or ".join(keyword_conditions)
        
        # Add username filter if provided
        if username:
            expr = f"({expr}) and username == '{username}'"
            
        print(f"[DEBUG] Keyword search expression: {expr}")
        
        try:
            # Execute query
            results = collection.query(
                expr=expr,
                output_fields=["id", "text", "username"],
                limit=top_k
            )
            
            print(f"[DEBUG] Keyword search returned {len(results)} chunks")
            
            # Sort results by relevance (count of keywords in each chunk)
            scored_results = []
            for result in results:
                # Double check username filter
                if username and result.get("username") != username:
                    continue
                    
                text = result.get("text", "").lower()
                # Count occurrences of each keyword
                score = sum(text.count(keyword) for keyword in keywords)
                scored_results.append({
                    "id": result.get("id"),
                    "text": result.get("text"),
                    "score": score
                })
            
            # Sort by score (highest first)
            scored_results.sort(key=lambda x: x["score"], reverse=True)
            
            # Extract texts from top results
            chunks = [result["text"] for result in scored_results[:top_k]]
            
            # Print samples from the first and last result for debugging
            if chunks:
                print(f"[DEBUG] First keyword result sample: {chunks[0][:100]}...")
                if len(chunks) > 1:
                    print(f"[DEBUG] Last keyword result sample: {chunks[-1][:100]}...")
                
            return chunks
        except Exception as query_error:
            print(f"[WARNING] Error in keyword query: {str(query_error)}")
            print("[INFO] Falling back to simpler query approach")
            
            # Fallback to a simpler approach with individual queries
            all_chunks = []
            for keyword in keywords:
                try:
                    # Add username filter to fallback query
                    keyword_expr = f"text like '%{keyword}%'"
                    if username:
                        keyword_expr += f" and username == '{username}'"
                        
                    keyword_results = collection.query(
                        expr=keyword_expr,
                        output_fields=["id", "text", "username"],
                        limit=min(10, top_k)  # Get fewer results per keyword
                    )
                    
                    # Double check username filter in results
                    for result in keyword_results:
                        if username and result.get("username") != username:
                            continue
                        all_chunks.append(result.get("text"))
                        
                except Exception as e:
                    print(f"[WARNING] Error querying for keyword '{keyword}': {str(e)}")
            
            # De-duplicate results
            seen = set()
            unique_chunks = []
            for chunk in all_chunks:
                chunk_hash = hash(chunk)
                if chunk_hash not in seen:
                    seen.add(chunk_hash)
                    unique_chunks.append(chunk)
            
            print(f"[DEBUG] Fallback keyword search returned {len(unique_chunks)} chunks")
            return unique_chunks[:top_k]
            
    except Exception as e:
        print(f"[ERROR] Exception in keyword_search: {str(e)}")
        return []

# Hybrid search combining vector and keyword approaches
@with_milvus_connection
def search_similar_chunks(query, top_k=10, hybrid_weight=0.7, use_advanced_search=True, username=None):
    """
    Advanced hybrid search combining vector similarity, keyword matching, and semantic analysis
    
    Args:
        query: The search query
        top_k: Number of results to return
        hybrid_weight: Weight for vector search (0.0-1.0), remaining weight goes to keyword search
        use_advanced_search: Whether to use advanced search techniques (False falls back to basic hybrid)
        username: The username to filter results by (only show results for this user)
    """
    print(f"[DEBUG] Starting enhanced search for query: '{query}'")
    print(f"[DEBUG] Parameters: top_k={top_k}, hybrid_weight={hybrid_weight}, use_advanced_search={use_advanced_search}, username={username}")
    
    try:
        # Step 1: Vector search 
        print(f"[DEBUG] Performing vector search with weight {hybrid_weight}")
        vector_top_k = min(top_k * 2, 30)  # Get more results for reranking
        
        # Encode query
        query_vec = embedder.encode([query])[0].tolist()
        print(f"[DEBUG] Query encoding complete, vector length: {len(query_vec)}")
        
        collection = Collection(COLLECTION_NAME)
        collection.load()
        
        # Execute vector search with username filter
        search_params = {
            "data": [query_vec],
            "anns_field": "embedding",
            "param": {"metric_type": "L2", "params": {"nprobe": 16}},
            "limit": vector_top_k,
            "output_fields": ["id", "text", "file_name", "username"]
        }
        
        if username:
            search_params["expr"] = f'username == "{username}"'
            
        vector_results = collection.search(**search_params)
        
        # Process vector results
        vector_chunks = []
        vector_ids = set()
        if vector_results and vector_results[0]:
            for hit in vector_results[0]:
                chunk_id = hit.entity.get("id")
                chunk_text = hit.entity.get("text")
                chunk_file = hit.entity.get("file_name", "")
                chunk_username = hit.entity.get("username", "")
                
                # Double check username filter
                if username and chunk_username != username:
                    continue
                    
                vector_distance = hit.distance
                
                # Lower distance is better for L2
                vector_score = 1.0 / (1.0 + vector_distance)
                
                vector_chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "file": chunk_file,
                    "score": vector_score,
                    "source": "vector"
                })
                vector_ids.add(chunk_id)
            
            print(f"[DEBUG] Vector search returned {len(vector_chunks)} chunks")
        
        # Step 2: Keyword search
        print(f"[DEBUG] Performing keyword search with weight {1.0 - hybrid_weight}")
        keyword_top_k = min(top_k * 2, 30)  # Get more results for reranking
        
        # Perform keyword search with username filter
        keyword_results = keyword_search(query, keyword_top_k, username)
        
        # Process keyword results
        keyword_chunks = []
        if keyword_results:
            # Get keyword match score for each result
            import re
            from collections import Counter
            
            # Clean query and extract keywords
            cleaned_query = re.sub(r'[^\w\s]', '', query.lower())
            words = cleaned_query.split()
            
            # Filter out common words (simple approach)
            common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'with', 'is', 'are', 'was', 'were'}
            keywords = [word for word in words if word not in common_words and len(word) > 2]
            
            for i, text in enumerate(keyword_results):
                # For keyword results, use position as part of score (better matches first)
                position_score = 1.0 - (i / len(keyword_results)) if len(keyword_results) > 1 else 1.0
                
                # Count keyword occurrences for scoring
                text_lower = text.lower()
                keyword_count = sum(text_lower.count(keyword) for keyword in keywords)
                keyword_density = keyword_count / (len(text_lower.split()) + 1)  # Add 1 to avoid division by zero
                
                # Get unique ID by hashing content (in case collection.query doesn't return IDs)
                chunk_id = hashlib.md5(text.encode()).hexdigest()
                
                # Skip if already in vector results
                if chunk_id in vector_ids:
                    continue
                
                # Compute a combined keyword score
                combined_keyword_score = (position_score * 0.4) + (keyword_density * 0.6)
                
                keyword_chunks.append({
                    "id": chunk_id,
                    "text": text,
                    "score": combined_keyword_score,
                    "source": "keyword" 
                })
            
            print(f"[DEBUG] Keyword search returned {len(keyword_chunks)} unique chunks")

        # Step 3: Advanced search techniques (if enabled)
        context_chunks = []
        if use_advanced_search and vector_chunks:
            print(f"[DEBUG] Performing advanced context-based search")
            
            # Find chunks that are contextually related to top vector results
            # by looking for chunks from the same files
            top_vector_files = set()
            for chunk in vector_chunks[:3]:  # Look at top 3 vector results
                file_name = chunk.get("file", "")
                if file_name:
                    top_vector_files.add(file_name)
            
            if top_vector_files:
                # Search for chunks from the same files
                try:
                    for file_name in top_vector_files:
                        # Add username filter to file query
                        expr = f'file_name == "{file_name}"'
                        if username:
                            expr += f' and username == "{username}"'
                            
                        context_results = collection.query(
                            expr=expr,
                            output_fields=["id", "text", "username"],
                            limit=5  # Limit per file
                        )
                        
                        for result in context_results:
                            chunk_id = result.get("id")
                            # Skip if already in other result sets
                            if chunk_id in vector_ids:
                                continue
                                
                            # Double check username filter
                            if username and result.get("username") != username:
                                continue
                                
                            chunk_text = result.get("text")
                            # Add with a moderate score
                            context_chunks.append({
                                "id": chunk_id,
                                "text": chunk_text,
                                "score": 0.5,  # Fixed score for context chunks
                                "source": "context"
                            })
                            
                    print(f"[DEBUG] Context search returned {len(context_chunks)} additional chunks")
                except Exception as context_error:
                    print(f"[WARNING] Error in context search: {str(context_error)}")
        
        # Step 4: Combine and rerank results
        print(f"[DEBUG] Combining and reranking results")
        combined_chunks = []
        
        # Add vector results with appropriate weight
        for chunk in vector_chunks:
            combined_chunks.append({
                "id": chunk["id"],
                "text": chunk["text"],
                "final_score": chunk["score"] * hybrid_weight,
                "source": chunk["source"]
            })
            
        # Add keyword results with appropriate weight
        for chunk in keyword_chunks:
            combined_chunks.append({
                "id": chunk["id"],
                "text": chunk["text"],
                "final_score": chunk["score"] * (1.0 - hybrid_weight),
                "source": chunk["source"]
            })
        
        # Add context chunks if advanced search is enabled
        if use_advanced_search and context_chunks:
            # Use a fixed weight for context chunks (lower than others)
            context_weight = 0.3
            for chunk in context_chunks:
                combined_chunks.append({
                    "id": chunk["id"],
                    "text": chunk["text"],
                    "final_score": chunk["score"] * context_weight,
                    "source": chunk["source"]
                })
            
        # Sort by final score
        combined_chunks.sort(key=lambda x: x["final_score"], reverse=True)
        
        # Remove duplicates and take top_k results
        seen_texts = set()
        unique_results = []
        
        for chunk in combined_chunks:
            # Use a hash of the first 100 chars to identify duplicates/near-duplicates
            text_hash = hash(chunk["text"][:100])
            if text_hash not in seen_texts:
                seen_texts.add(text_hash)
                unique_results.append(chunk)
                
                # Stop once we have enough results
                if len(unique_results) >= top_k:
                    break
        
        # Extract texts from top results
        result_chunks = [chunk["text"] for chunk in unique_results]
        
        print(f"[DEBUG] Final search returned {len(result_chunks)} unique chunks")
        source_distribution = Counter([chunk["source"] for chunk in unique_results])
        print(f"[DEBUG] Source distribution in top results: {dict(source_distribution)}")
            
        # Print samples from the first and last result for debugging
        if result_chunks:
            print(f"[DEBUG] First result sample: {result_chunks[0][:100]}...")
            if len(result_chunks) > 1:
                print(f"[DEBUG] Last result sample: {result_chunks[-1][:100]}...")
        
        return result_chunks
    except Exception as e:
        print(f"[ERROR] Exception in search_similar_chunks: {str(e)}")
        # In case of error, fall back to standard vector search
        try:
            print(f"[DEBUG] Falling back to standard vector search")
            query_vec = embedder.encode([query])[0].tolist()
            collection = Collection(COLLECTION_NAME)
            collection.load()
            
            # Add username filter to fallback search
            search_params = {
                "data": [query_vec],
                "anns_field": "embedding",
                "param": {"metric_type": "L2", "params": {"nprobe": 10}},
                "limit": top_k,
                "output_fields": ["text", "username"]
            }
            
            if username:
                search_params["expr"] = f'username == "{username}"'
                
            results = collection.search(**search_params)
            
            if results and results[0]:
                # Double check username filter in results
                chunks = []
                for hit in results[0]:
                    if username and hit.entity.get("username") != username:
                        continue
                    chunks.append(hit.entity.get("text"))
                    
                print(f"[DEBUG] Fallback search returned {len(chunks)} chunks")
                return chunks
            else:
                return []
        except Exception as fallback_error:
            print(f"[ERROR] Fallback search also failed: {str(fallback_error)}")
            return []

# Get conversation history
def get_conversation_history(conversation_id):
    """Retrieve all messages from the conversation history"""
    print(f"[DEBUG] Retrieving conversation history for conversation ID: {conversation_id}")
    try:
        conversation_doc = conversations.find_one({"_id": ObjectId(conversation_id)})
        if conversation_doc and "messages" in conversation_doc:
            # Get all messages, no limit
            all_messages = conversation_doc["messages"]
            history_text = ""
            for msg in all_messages:
                role = msg["role"]
                message = msg["message"]
                history_text += f"{role.capitalize()}: {message}\n\n"
            return history_text
        return ""
    except Exception as e:
        print(f"[ERROR] Exception retrieving conversation history: {str(e)}")
        return ""

def get_all_vector_entries():
    """Retrieve all vector entries without pagination"""
    print(f"[DEBUG] Retrieving all vector entries")
    try:
        # Connect to collection
        collection = Collection(COLLECTION_NAME)
        collection.load()
        
        # First, get all unique file hashes to identify the files
        print(f"[DEBUG] Retrieving unique file hashes")
        file_hash_results = collection.query(
            expr=f'username == "{session["username"]}"',  # Only get user's files
            output_fields=["file_hash", "file_name"],
            limit=10000  # High enough to get all unique files
        )
        
        # Extract unique file hashes
        unique_files = {}
        for result in file_hash_results:
            file_hash = result.get("file_hash", "")
            if file_hash and file_hash not in unique_files:
                unique_files[file_hash] = result.get("file_name", "unknown")
        
        print(f"[DEBUG] Found {len(unique_files)} unique files")
        
        # Initialize return structure
        files_data = {}
        total_count = 0
        
        # For each file hash, get all entries
        for file_hash, file_name in unique_files.items():
            print(f"[DEBUG] Retrieving entries for file: {file_name} (hash: {file_hash})")
            all_file_entries = []
            
            # Retrieve in chunks to handle Milvus limitations
            offset = 0
            while True:
                file_entries = collection.query(
                    expr=f'file_hash == "{file_hash}" and username == "{session["username"]}"',  # Only get user's entries
                    output_fields=["id", "text", "file_hash", "file_name"],
                    offset=offset,
                    limit=10000  # Chunk size
                )
                
                if not file_entries:
                    break
                    
                all_file_entries.extend(file_entries)
                file_count = len(file_entries)
                print(f"[DEBUG] Retrieved {file_count} entries for file {file_name}")
                
                # If we received fewer results than the limit, we've reached the end
                if file_count < 10000:
                    break
                    
                # Move to next chunk
                offset += 10000
            
            # Process entries for this file
            file_entry_data = []
            for entry in all_file_entries:
                chunk_text = entry.get("text", "")
                # Truncate text preview
                preview = chunk_text[:100] + "..." if len(chunk_text) > 100 else chunk_text
                
                file_entry_data.append({
                    "id": entry.get("id"),
                    "preview": preview,
                    "full_text": chunk_text
                })
            
            # Add file data to the collection
            if file_entry_data:
                entry_count = len(file_entry_data)
                files_data[file_hash] = {
                    "file_name": file_name,
                    "entries": file_entry_data,
                    "count": entry_count
                }
                total_count += entry_count
        
        # Convert to list for template
        files_list = []
        for file_hash, data in files_data.items():
            files_list.append({
                "file_hash": file_hash,
                "file_name": data["file_name"],
                "entries": data["entries"],
                "count": data["count"]
            })
        
        # Sort files by name for better display
        files_list.sort(key=lambda x: x["file_name"])
        
        print(f"[DEBUG] Retrieved a total of {total_count} entries across {len(files_list)} files")
        
        return {
            "files": files_list,
            "total_count": total_count
        }
    except Exception as e:
        print(f"[ERROR] Exception retrieving vector entries: {str(e)}")
        return {"files": [], "total_count": 0}

def ask_llm_via_chatgpt(prompt):
    """Send a prompt to ChatGPT API and return the response"""
    print(f"[INFO] Preparing request for ChatGPT API")
    
    # If API key is not set, return an error message
    if not OPENAI_API_KEY:
        print("[WARNING] No OpenAI API key set.")
        return "I'm sorry, but the API key for ChatGPT is not configured. Please set up your OpenAI API key."
    
    # Set up headers for authentication
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "gpt-3.5-turbo",  # Most cost-effective model
        "messages": [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 1000,  # Reduced from 2000 to save tokens
        "presence_penalty": 0.1,  # Slight penalty for repetition
        "frequency_penalty": 0.1  # Slight penalty for frequent tokens
    }
    
    try:
        print(f"[INFO] Sending request to ChatGPT API")
        
        # Use OpenAI's API endpoint
        api_url = "https://api.openai.com/v1/chat/completions"
        
        # Add retry logic
        max_retries = 3
        retry_delay = 2  # seconds
        
        for attempt in range(max_retries):
            try:
                res = requests.post(
                    api_url, 
                    headers=headers, 
                    json=payload, 
                    timeout=30
                )
                
                # If successful, break the retry loop
                if res.status_code == 200:
                    break
                    
                # If rate limited, wait and retry
                if res.status_code == 429:
                    if attempt < max_retries - 1:  # Don't wait on last attempt
                        wait_time = retry_delay * (attempt + 1)  # Exponential backoff
                        print(f"[INFO] Rate limited. Waiting {wait_time} seconds before retry...")
                        time.sleep(wait_time)
                        continue
                
                # For other errors, try to parse the error message
                error_message = f"Status code: {res.status_code}"
                try:
                    error_json = res.json()
                    if "error" in error_json:
                        error_msg = error_json["error"]
                        if isinstance(error_msg, dict) and "message" in error_msg:
                            error_msg = error_msg["message"]
                        error_message = f"{error_message} - {error_msg}"
                except:
                    pass
                
                # If this was the last attempt, return the error
                if attempt == max_retries - 1:
                    return f"Error from API: {error_message}"
                
            except requests.exceptions.Timeout:
                if attempt == max_retries - 1:
                    print("[ERROR] Request timed out after all retries")
                    return "Sorry, the request timed out after multiple attempts. Please try again later."
                print(f"[WARNING] Request timed out, attempt {attempt + 1} of {max_retries}")
                time.sleep(retry_delay)
                
            except requests.exceptions.ConnectionError:
                if attempt == max_retries - 1:
                    print("[ERROR] Connection error after all retries")
                    return "Sorry, there was a connection error. Please check your internet connection and try again."
                print(f"[WARNING] Connection error, attempt {attempt + 1} of {max_retries}")
                time.sleep(retry_delay)
                
            except requests.exceptions.RequestException as e:
                if attempt == max_retries - 1:
                    print(f"[ERROR] Request failed after all retries: {str(e)}")
                    return f"Sorry, there was an error communicating with the server: {str(e)}"
                print(f"[WARNING] Request failed, attempt {attempt + 1} of {max_retries}: {str(e)}")
                time.sleep(retry_delay)
        
        print(f"[INFO] API response status code: {res.status_code}")
        
        if res.status_code == 200:
            try:
                response_json = res.json()
                
                # Check for required fields
                if "choices" not in response_json:
                    print(f"[ERROR] 'choices' key not found in response")
                    return "Error: API response missing expected 'choices' field. Please check your API key and model configuration."
                
                if not response_json["choices"] or len(response_json["choices"]) == 0:
                    print(f"[ERROR] Empty choices array in response")
                    return "Error: API returned empty choices. Please try again or check your API configuration."
                
                # Check if the structure is as expected
                if "message" not in response_json["choices"][0]:
                    return "Error: Unexpected API response format. Please check the API documentation."
                
                if "content" not in response_json["choices"][0]["message"]:
                    return "Error: API response missing content field. Please try again."
                
                content = response_json["choices"][0]["message"]["content"]
                return content
            except KeyError as e:
                print(f"[ERROR] KeyError parsing response: {str(e)}")
                return f"Error parsing API response: {str(e)}. Please check your API configuration."
            except Exception as e:
                print(f"[ERROR] Exception parsing response: {str(e)}")
                return f"Error processing API response: {str(e)}"
        else:
            print(f"[ERROR] ChatGPT API returned error: {res.status_code}")
            
            # Try to parse error message from JSON if possible
            error_message = f"Status code: {res.status_code}"
            try:
                error_json = res.json()
                if "error" in error_json:
                    error_msg = error_json["error"]
                    if isinstance(error_msg, dict) and "message" in error_msg:
                        error_msg = error_msg["message"]
                    error_message = f"{error_message} - {error_msg}"
            except:
                pass
                
            return f"Error from API: {error_message}"
            
    except requests.exceptions.RequestException as e:
        print(f"[ERROR] Network error: {str(e)}")
        return f"Network error connecting to API: {str(e)}"
    except Exception as e:
        print(f"[ERROR] Exception: {str(e)}")
        return f"Error making API request: {str(e)}"

# Auth routes
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        user = users.find_one({"username": username})
        if user and check_password_hash(user["password"], password):
            session["username"] = username
            return redirect(url_for("index"))
        return render_template("login.html", error="Invalid username or password")
    return render_template("login.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        if users.find_one({"username": username}):
            return render_template("register.html", error="Username already exists")
        users.insert_one({
            "username": username,
            "password": generate_password_hash(password)
        })
        return redirect(url_for("login"))
    return render_template("register.html")

@app.route("/logout")
def logout():
    session.pop("username", None)
    session.pop("conversation_id", None)
    return redirect(url_for("login"))

@app.route("/upload_progress", methods=["POST"])
@with_milvus_connection
def upload_progress():
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401
        
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
        
    files = request.files.getlist("file")
    all_chunks = []
    processed_files = []
    duplicate_files = []
    username = session["username"]  # Get username once at the start
    
    # First, collect all file hashes from current upload
    file_hashes = {}
    for file in files:
        if file.filename:
            file_hash = calculate_file_hash(file)
            if file_hash in file_hashes:
                duplicate_files.append(file.filename)
            else:
                file_hashes[file_hash] = file.filename
    
    # Check against existing files in Milvus
    collection = Collection(COLLECTION_NAME)
    collection.load()
    
    try:
        # Get all existing file hashes for the current user only
        existing_hashes = set()
        expr = f"file_hash != '' and username == '{username}'"  # Only check for this user's files
        batch_size = 1000  # Smaller batch size to stay within limits
        offset = 0
        
        while True:
            results = collection.query(
                expr=expr,
                output_fields=["file_hash", "file_name"],
                offset=offset,
                limit=batch_size
            )
            
            if not results:
                break
                
            for result in results:
                existing_hashes.add(result["file_hash"])
            
            if len(results) < batch_size:
                break
                
            offset += batch_size
        
        # Check current files against existing hashes and process non-duplicates
        for file in files:
            if file.filename:
                file_hash = calculate_file_hash(file)
                if file_hash in existing_hashes:
                    duplicate_files.append(file.filename)
                else:
                    processed_files.append(file.filename)
                    text = extract_text_from_file(file)
                    if text:
                        chunks = chunk_text(text)
                        for chunk in chunks:
                            chunk_id = str(uuid.uuid4())
                            data = {
                                "id": chunk_id,
                                "text": chunk,
                                "file_hash": file_hash,
                                "file_name": file.filename,
                                "username": username
                            }
                            all_chunks.append(data)
        
        if all_chunks:
            try:
                # Create a conversation for this upload if needed
                conversation_id = None
                is_new_conversation = False
                
                if "conversation_id" in session:
                    try:
                        conversation_id = ObjectId(session["conversation_id"])
                        conversation_doc = conversations.find_one({"_id": conversation_id})
                        if not conversation_doc:
                            # Invalid conversation ID
                            session.pop("conversation_id", None)
                            conversation_id = None
                    except:
                        session.pop("conversation_id", None)
                
                # Create a new conversation for these files if needed
                if conversation_id is None:
                    # Create a new conversation
                    result = conversations.insert_one({
                        "username": username,
                        "title": f"Document Upload: {', '.join(processed_files[:2])}" + ("..." if len(processed_files) > 2 else ""),
                        "messages": [],
                        "created_at": datetime.now()
                    })
                    conversation_id = result.inserted_id
                    session["conversation_id"] = str(conversation_id)
                    is_new_conversation = True
                
                def generate():
                    for progress in store_chunks_in_milvus(all_chunks, username):  # Pass username here
                        yield progress
                    
                    # Add system message about upload to the conversation
                    if conversation_id:
                        upload_message = f"Files uploaded successfully. {len(all_chunks)} chunks extracted and processed from {len(processed_files)} file(s)."
                        conversations.update_one(
                            {"_id": conversation_id},
                            {"$push": {
                                "messages": {
                                    "role": "assistant", 
                                    "message": upload_message, 
                                    "timestamp": datetime.now()
                                }
                            }}
                        )
                
                return Response(generate(), mimetype='text/event-stream')
            except Exception as e:
                return jsonify({"error": str(e)}), 500
        else:
            if duplicate_files:
                return jsonify({
                    "error": "All files have already been uploaded before do not upload again",
                    "duplicates": duplicate_files
                }), 400
            else:
                return jsonify({"error": "No text extracted from files"}), 400
            
    except Exception as e:
        print(f"[ERROR] Error checking for duplicates: {str(e)}")
        return jsonify({"error": "Error checking for duplicate files"}), 500

@app.route("/conversation/<conversation_id>", methods=["GET"])
def get_conversation(conversation_id):
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401
        
    try:
        # Set the session conversation ID
        session["conversation_id"] = conversation_id
        
        # Find the conversation
        conversation = conversations.find_one({
            "_id": ObjectId(conversation_id),
            "username": session["username"]
        })
        
        if not conversation:
            return jsonify({"error": "Conversation not found"}), 404
        
        # Log the number of messages for debugging
        messages = conversation.get("messages", [])
        print(f"[DEBUG] Returning conversation with {len(messages)} messages")
            
        # Return all messages, no limit
        return jsonify({
            "messages": messages,
            "title": conversation.get("title", "Untitled Conversation")
        })
    except Exception as e:
        print(f"[ERROR] Exception in get_conversation: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route("/update_title", methods=["POST"])
def update_title():
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json()
    new_title = data.get("title")
    conversation_id = data.get("conversation_id") or session.get("conversation_id")

    if not conversation_id:
        return jsonify({"error": "No active conversation"}), 400

    if not new_title:
        return jsonify({"error": "Title cannot be empty"}), 400

    conversations.update_one(
        {"_id": ObjectId(conversation_id), "username": session["username"]},
        {"$set": {"title": new_title}}
    )

    return jsonify({"success": True})

@app.route("/delete_conversation", methods=["POST"])
def delete_conversation():
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json(silent=True) or {}
    conversation_id = data.get("conversation_id") or session.get("conversation_id")

    if not conversation_id:
        return jsonify({"error": "No active conversation"}), 400

    conversations.delete_one({
        "_id": ObjectId(conversation_id),
        "username": session["username"]
    })

    # If the deleted conversation was the active one, clear it from session
    if session.get("conversation_id") == conversation_id:
        session.pop("conversation_id", None)

    return jsonify({"success": True})

@app.route("/manage_vectors")
def manage_vectors():
    if "username" not in session:
        return redirect(url_for("login"))
    
    # Get vector entries
    vector_data = get_all_vector_entries()
    
    return render_template(
        "manage_vectors.html",
        username=session["username"],
        files=vector_data["files"],
        total_count=vector_data["total_count"]
    )

@app.route("/delete_vector_entry", methods=["POST"])
def delete_vector_entry():
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401
    
    data = request.get_json(silent=True) or {}
    entry_id = data.get("entry_id")
    
    if not entry_id:
        return jsonify({"error": "No entry ID provided"}), 400
    
    try:
        # Connect to collection
        collection = Collection(COLLECTION_NAME)
        collection.load()
        
        # Delete entry
        expr = f'id == "{entry_id}"'
        collection.delete(expr)
        
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERROR] Exception deleting vector entry: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route("/delete_vector_file", methods=["POST"])
def delete_vector_file():
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401
    
    data = request.get_json(silent=True) or {}
    file_hash = data.get("file_hash")
    
    if not file_hash:
        return jsonify({"error": "No file hash provided"}), 400
    
    try:
        # Connect to collection
        collection = Collection(COLLECTION_NAME)
        collection.load()
        
        # Delete all entries with the given file_hash
        expr = f'file_hash == "{file_hash}"'
        collection.delete(expr)
        
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERROR] Exception deleting vector file: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route("/new_conversation_id", methods=["POST"])
def new_conversation_id():
    """Create a new empty conversation and return its ID"""
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401

    # Clear any existing conversation_id from session
    if "conversation_id" in session:
        session.pop("conversation_id", None)
    
    # Create a new empty conversation
    result = conversations.insert_one({
        "username": session["username"],
        "title": "New Conversation",
        "messages": [],
        "created_at": datetime.now()
    })
    
    # Store new conversation ID in session
    new_id = str(result.inserted_id)
    session["conversation_id"] = new_id
    
    return jsonify({
        "success": True,
        "conversation_id": new_id
    })

@app.route("/query", methods=["POST"])
def handle_query():
    """Handle AJAX query requests"""
    if "username" not in session:
        return jsonify({"error": "Not authenticated"}), 401

    # Process the query first before creating any conversation
    query = request.form.get("query", "").strip()
    if not query:
        return jsonify({"error": "Empty query"}), 400

    # Get hybrid weight parameter (default to 0.7 if not provided)
    try:
        hybrid_weight = float(request.form.get("hybrid_weight", 0.7))
        # Ensure it's within valid range
        hybrid_weight = min(max(hybrid_weight, 0.0), 1.0)
    except (ValueError, TypeError):
        hybrid_weight = 0.7

    print(f"[DEBUG] Processing query: '{query}' with hybrid_weight: {hybrid_weight}")
    print(f"[DEBUG] Request form data: {dict(request.form)}")
    print(f"[DEBUG] Session data: {dict(session)}")
    
    conversation_id = None
    conversation_doc = None
    is_new_conversation = False
    
    # Check if client sent a conversation_id in the form data
    client_conversation_id = request.form.get("conversation_id")
    if client_conversation_id:
        print(f"[DEBUG] Client sent conversation_id in form data: {client_conversation_id}")
        try:
            # Use this ID instead of the session one if it exists
            conversation_id = ObjectId(client_conversation_id)
            conversation_doc = conversations.find_one({"_id": conversation_id})
            if conversation_doc:
                session["conversation_id"] = client_conversation_id
                print(f"[DEBUG] Using client conversation_id: {client_conversation_id}")
            else:
                print(f"[DEBUG] Client conversation_id not found in database")
                conversation_id = None
        except Exception as e:
            print(f"[DEBUG] Error using client conversation_id: {str(e)}")
            conversation_id = None
    # If no valid client conversation ID, check session
    elif "conversation_id" in session:
        try:
            print(f"[DEBUG] Found conversation_id in session: {session['conversation_id']}")
            conversation_id = ObjectId(session["conversation_id"])
            conversation_doc = conversations.find_one({"_id": conversation_id})
            if not conversation_doc:
                print(f"[DEBUG] Conversation {conversation_id} not found in database")
                # Conversation was deleted or doesn't exist
                session.pop("conversation_id", None)
                conversation_id = None
            else:
                print(f"[DEBUG] Found conversation document with {len(conversation_doc.get('messages', []))} messages")
        except Exception as e:
            print(f"[DEBUG] Exception checking conversation: {str(e)}")
            session.pop("conversation_id", None)
            conversation_id = None
    else:
        print(f"[DEBUG] No conversation_id in session or form data")
    
    # Check if we need to create a new conversation
    if conversation_id is None:
        print(f"[DEBUG] Creating new conversation")
        # Create a new conversation - we have a query now so it's a good time
        result = conversations.insert_one({
            "username": session["username"],
            "title": "Untitled Conversation",
            "messages": [],
            "created_at": datetime.now()
        })
        conversation_id = result.inserted_id
        print(f"[DEBUG] Created new conversation with ID: {conversation_id}")
        session["conversation_id"] = str(conversation_id)
        conversation_doc = conversations.find_one({"_id": conversation_id})
        is_new_conversation = True
        print(f"[DEBUG] is_new_conversation set to: {is_new_conversation}")
    
    try:
        # Use hybrid search with the specified weight
        top_chunks = search_similar_chunks(query, top_k=10, hybrid_weight=hybrid_weight, username=session["username"])
        document_context = "\n".join(top_chunks)
        
        # Get conversation history
        conversation_history = get_conversation_history(str(conversation_id))
        
        prompt = (
            "You are Bio-Logic, an advanced biomedical research assistant trained to understand, summarize, explain, and guide based on complex medical documents.\n"
            "You will determine the intent of the user's question and respond accordingly:\n\n"
            "— If the user is asking for a summary, provide a clear and concise summary of the relevant information.\n"
            "— If the user is asking for an explanation, define and clarify biomedical terms or findings using the context.\n"
            "— If the user is seeking clinical or research-based guidance, provide evidence-based insights strictly from the documents.\n\n"
            "FORMAT YOUR RESPONSE FOR READABILITY:\n"
            "1. Use clear section headers (bold with **) when presenting different topics\n"
            "2. Use bullet points (•) for listing items within a section\n"
            "3. Use numbered lists (1., 2., 3.) for sequential steps or prioritized information\n"
            "4. Emphasize key terms or findings with italics (*term*)\n"
            "5. Break your response into logical paragraphs with appropriate spacing\n"
            "6. For data or statistics, present them in a structured format with clear labels\n\n"
            "Do not speculate or fabricate. If the answer is not in the documents, respond with:\n"
            "'I don't have that information in the documents you've provided.'\n\n"
            "Maintain a clear, professional tone appropriate for medical researchers and clinicians.\n\n"
            f"--- Document Context ---\n{document_context}\n\n"
            f"--- Previous Conversation ---\n{conversation_history}\n\n"
            f"--- User Question ---\n{query}\n\n"
            "Answer:"
        )
        answer = ask_llm_via_chatgpt(prompt)
        
        # Append messages to conversation
        conversations.update_one(
            {"_id": conversation_id},
            {"$push": {
                "messages": {
                    "$each": [
                        {"role": "user", "message": query, "timestamp": datetime.now()},
                        {"role": "assistant", "message": answer, "timestamp": datetime.now()}
                    ]
                }
            }}
        )
        
        # Generate/update conversation title if needed
        updated_conversation = conversations.find_one({"_id": conversation_id})
        conversations_history = updated_conversation.get("messages", [])
        
        # Update conversation title after a few messages or immediately for first message
        # if this is a new conversation
        if len(conversations_history) >= 3 or is_new_conversation:
            # Extract text from messages
            conversation_text = ""
            for msg in conversations_history:
                if msg["role"] == "user":
                    conversation_text += msg["message"] + " "
            
            # Generate title using the LLM
            if conversation_text:
                title_prompt = f"Generate a short, descriptive title (max 5 words) for this conversation: {conversation_text[:500]}"
                conversation_title = ask_llm_via_chatgpt(title_prompt)
                conversations.update_one(
                    {"_id": conversation_id},
                    {"$set": {"title": conversation_title}}
                )
                updated_conversation = conversations.find_one({"_id": conversation_id})
                
        # Return JSON response with the answer and updated title
        response_data = {
            "answer": answer,
            "title": updated_conversation.get("title", "Untitled Conversation"),
            "conversation_id": str(conversation_id),
            "is_new_conversation": is_new_conversation,
            "hybrid_weight_used": hybrid_weight
        }
        print(f"[DEBUG] Returning response with is_new_conversation: {is_new_conversation}")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"[ERROR] Error processing query: {str(e)}")
        return jsonify({"error": f"Error processing your query: {str(e)}"}), 500

# Main route
@app.route("/", methods=["GET", "POST"])
def index():
    if "username" not in session:
        return redirect(url_for("login"))

    answer, uploaded_file, upload_time = "", "", ""
    conversation_title = "New Conversation"
    conversation_doc = None
    conversations_history = []
    conversation_id = None

    # Only retrieve the conversation if there's a conversation_id in the session
    if "conversation_id" in session:
        try:
            conversation_id = ObjectId(session["conversation_id"])
            conversation_doc = conversations.find_one({"_id": conversation_id})
            if conversation_doc:
                # Get ALL messages from the conversation, no limit
                conversations_history = conversation_doc.get("messages", [])
                conversation_title = conversation_doc.get("title", "Untitled Conversation")
                
                # Generate title only if not already set or if it's the default
                if conversation_title == "Untitled Conversation" and conversations_history:
                    # Get messages from the middle third of the conversation
                    mid_start = len(conversations_history) // 3
                    mid_end = 2 * len(conversations_history) // 3
                    mid_messages = conversations_history[mid_start:mid_end]
                    
                    # Extract text from messages
                    conversation_text = ""
                    for msg in mid_messages:
                        if msg["role"] == "user":
                            conversation_text += msg["message"] + " "
                    
                    # Generate title using the LLM
                    if conversation_text:
                        title_prompt = f"Generate a short, descriptive title (max 5 words) for this conversation: {conversation_text[:500]}"
                        conversation_title = ask_llm_via_chatgpt(title_prompt)
                        # Update the conversation title in the database
                        conversations.update_one(
                            {"_id": conversation_id},
                            {"$set": {"title": conversation_title}}
                        )
        except:
            session.pop("conversation_id", None)

    if request.method == "POST":
        # Check if we need to create a new conversation
        if conversation_id is None and ("file" in request.files or "query" in request.form):
            # Create a new conversation only if there's an actual file upload or query
            has_files = False
            query = ""
            
            if "file" in request.files:
                file = request.files["file"]
                if file.filename:
                    has_files = True
            
            if "query" in request.form:
                query = request.form["query"].strip()
            
            if has_files or query:
                # Create a new conversation since we have actual content
                result = conversations.insert_one({
                    "username": session["username"],
                    "title": "Untitled Conversation",
                    "messages": [],
                    "created_at": datetime.now()
                })
                conversation_id = result.inserted_id
                session["conversation_id"] = str(conversation_id)
                conversation_doc = conversations.find_one({"_id": conversation_id})
                conversations_history = []

        if "file" in request.files:
            files = request.files.getlist("file")
            all_chunks = []
            processed_files = []
            
            for file in files:
                if file.filename:
                    processed_files.append(file.filename)
                    text = extract_text_from_file(file)
                    if text:
                        chunks = chunk_text(text)
                        all_chunks.extend(chunks)
            
            if all_chunks:
                upload_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                try:
                    store_chunks_in_milvus(all_chunks, session["username"])
                    uploaded_file = ", ".join(processed_files)
                    answer = f"Uploaded and processed {len(all_chunks)} chunks from {len(processed_files)} file(s)."
                    
                    # Add system message about upload
                    if conversation_id:
                        conversations.update_one(
                            {"_id": conversation_id},
                            {"$push": {
                                "messages": {
                                    "role": "assistant", 
                                    "message": f"Files uploaded successfully. {len(all_chunks)} chunks extracted and processed from {len(processed_files)} file(s).", 
                                    "timestamp": datetime.now()
                                }
                            }}
                        )
                        conversation_doc = conversations.find_one({"_id": conversation_id})
                        conversations_history = conversation_doc.get("messages", [])
                except Exception as e:
                    answer = f"Error processing files: {str(e)}"
            else:
                answer = "No extractable text found in the uploaded files."

        elif "query" in request.form:
            query = request.form["query"].strip()
            if query:
                try:
                    # Get hybrid weight parameter (default to 0.7 if not provided)
                    try:
                        hybrid_weight = float(request.form.get("hybrid_weight", 0.7))
                        # Ensure it's within valid range
                        hybrid_weight = min(max(hybrid_weight, 0.0), 1.0)
                    except (ValueError, TypeError):
                        hybrid_weight = 0.7
                        
                    print(f"[DEBUG] Using hybrid weight: {hybrid_weight} for query: {query}")
                    
                    # Use the hybrid search with specified weight
                    top_chunks = search_similar_chunks(query, top_k=10, hybrid_weight=hybrid_weight, username=session["username"])
                    document_context = "\n".join(top_chunks)
                    
                    # Get conversation history
                    conversation_history = get_conversation_history(str(conversation_id))
                    
                    prompt = (
                        "You are Bio-Logic, an advanced biomedical research assistant trained to understand, summarize, explain, and guide based on complex medical documents.\n"
                        "You will determine the intent of the user's question and respond accordingly:\n\n"
                        "— If the user is asking for a summary, provide a clear and concise summary of the relevant information.\n"
                        "— If the user is asking for an explanation, define and clarify biomedical terms or findings using the context.\n"
                        "— If the user is seeking clinical or research-based guidance, provide evidence-based insights strictly from the documents.\n\n"
                        "FORMAT YOUR RESPONSE FOR READABILITY:\n"
                        "1. Use clear section headers (bold with **) when presenting different topics\n"
                        "2. Use bullet points (•) for listing items within a section\n"
                        "3. Use numbered lists (1., 2., 3.) for sequential steps or prioritized information\n"
                        "4. Emphasize key terms or findings with italics (*term*)\n"
                        "5. Break your response into logical paragraphs with appropriate spacing\n"
                        "6. For data or statistics, present them in a structured format with clear labels\n\n"
                        "Do not speculate or fabricate. If the answer is not in the documents, respond with:\n"
                        "'I don't have that information in the documents you've provided.'\n\n"
                        "Maintain a clear, professional tone appropriate for medical researchers and clinicians.\n\n"
                        f"--- Document Context ---\n{document_context}\n\n"
                        f"--- Previous Conversation ---\n{conversation_history}\n\n"
                        f"--- User Question ---\n{query}\n\n"
                        "Answer:"
                    )
                    answer = ask_llm_via_chatgpt(prompt)
                    
                    # Append messages to conversation
                    if conversation_id:
                        conversations.update_one(
                            {"_id": conversation_id},
                            {"$push": {
                                "messages": {
                                    "$each": [
                                        {"role": "user", "message": query, "timestamp": datetime.now()},
                                        {"role": "assistant", "message": answer, "timestamp": datetime.now()}
                                    ]
                                }
                            }}
                        )
                        
                        # Update conversation title after a few messages
                        if len(conversations_history) >= 3:
                            # Get messages from the middle third
                            mid_start = len(conversations_history) // 3
                            mid_end = 2 * len(conversations_history) // 3
                            mid_messages = conversations_history[mid_start:mid_end]
                            
                            # Extract text from messages
                            conversation_text = ""
                            for msg in mid_messages:
                                if msg["role"] == "user":
                                    conversation_text += msg["message"] + " "
                            
                            # Generate title using the LLM
                            if conversation_text:
                                title_prompt = f"Generate a short, descriptive title (max 5 words) for this conversation: {conversation_text[:500]}"
                                conversation_title = ask_llm_via_chatgpt(title_prompt)
                                conversations.update_one(
                                    {"_id": conversation_id},
                                    {"$set": {"title": conversation_title}}
                                )
                        
                        conversation_doc = conversations.find_one({"_id": conversation_id})
                        conversations_history = conversation_doc.get("messages", [])
                        conversation_title = conversation_doc.get("title", "Untitled Conversation")
                except Exception as e:
                    answer = f"Error processing your query: {str(e)}"

    # Get all user conversations for the sidebar
    all_user_conversations = list(conversations.find(
        {"username": session["username"]},
        {"title": 1, "created_at": 1}
    ).sort("created_at", -1))

    # Get default hybrid search settings
    default_hybrid_weight = 0.7

    return render_template(
        "index.html", 
        uploaded_file=uploaded_file,
        upload_time=upload_time, 
        answer=answer, 
        conversations=conversations_history,
        conversation_title=conversation_title,
        all_user_conversations=all_user_conversations,
        username=session["username"],
        default_hybrid_weight=default_hybrid_weight
    )

if __name__ == "__main__":
    # Print status message
    print("\n===== Medical Document Assistant =====")
    print("Server starting at http://localhost:5000\n")
    
    if not OPENAI_API_KEY:
        print("WARNING: No OpenAI API key found. Please set your OPENAI_API_KEY in the .env file.")
    else:
        print("INFO: OpenAI API key configured. Full functionality available.\n")
        
    app.run(debug=True,host="0.0.0.0",port=5000)